#!/usr/bin/env python3
"""Generate a CN patent disclosure .docx from merged markdown.

Usage:
  python scripts/generate_docx.py <input.md> <output.docx>

Dependencies:
  python-docx (already in requirements.txt)

Features:
  - Sets CN patent standard fonts (SimSun body, SimHei headings)
  - Embed images referenced in markdown (![alt](path))
  - Convert Mermaid code blocks to [插图] placeholder (user replaces manually)
  - Auto-size page for A4
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def generate_docx(md_path: str, docx_path: str) -> bool:
    md_file = Path(md_path)
    if not md_file.exists():
        print(f"ERROR: {md_path} not found")
        return False

    doc = Document()

    # ── Page setup: A4 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    content = md_file.read_text(encoding="utf-8")
    lines = content.split("\n")
    base_dir = md_file.parent

    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading: # 一、技术领域
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            _set_heading_font(p, 'SimHei', Pt(16))
            i += 1
            continue

        # Heading: ## 2.1 xxx
        if line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            _set_heading_font(p, 'SimHei', Pt(14))
            i += 1
            continue

        # Heading: ### xxx
        if line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            _set_heading_font(p, 'SimHei', Pt(13))
            i += 1
            continue

        # Image: ![alt](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)
            img_path = base_dir / img_rel_path
            if img_path.exists():
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))
                    caption = doc.add_paragraph(f"图：{alt_text}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_run_font(caption.runs[0], 'SimSun', Pt(10))
                except Exception as e:
                    doc.add_paragraph(f"[插图：{alt_text} - 加载失败: {e}]")
            else:
                doc.add_paragraph(f"[插图：{alt_text} - 文件未找到: {img_rel_path}]")
            i += 1
            continue

        # Mermaid code block: skip ```mermaid ... ```, insert placeholder
        if line.strip().startswith("```mermaid"):
            doc.add_paragraph("[Mermaid 流程图 - 请手动替换为渲染后的图片]")
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1  # skip closing ```
            continue

        # Quote: >
        if line.startswith("> "):
            text = line[2:].strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Cm(1)
                _set_run_font(p.runs[0] if p.runs else p.add_run(text), 'SimSun', Pt(10))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        # Merge consecutive lines into one paragraph
        para_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not lines[j].startswith("#") and not lines[j].startswith("```") and not lines[j].startswith(">") and not re.match(r'!\[', lines[j]):
            para_lines.append(lines[j])
            j += 1

        text = "".join(para_lines)
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符缩进
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.name = 'SimSun'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        i = j

    doc.save(docx_path)
    return Path(docx_path).exists()


def _set_heading_font(paragraph, font_name: str, size: Pt):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = size
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _set_run_font(run, font_name: str, size: Pt):
    run.font.name = font_name
    run.font.size = size
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python scripts/generate_docx.py <input.md> <output.docx>")
        sys.exit(1)
    ok = generate_docx(sys.argv[1], sys.argv[2])
    sys.exit(0 if ok else 1)
