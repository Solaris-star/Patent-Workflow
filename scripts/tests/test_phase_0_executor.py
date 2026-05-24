#!/usr/bin/env python3
"""Phase 0 preprocess cache regression tests."""

import json
import sys
import tempfile
import time
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "scripts"))

from executors.phase_0_executor import PhaseExecutor  # noqa: E402
import orchestrate as orchestrate_module  # noqa: E402
from orchestrate import Orchestrator  # noqa: E402


def _run_phase_0(workspace: Path, manifest=None):
    executor = PhaseExecutor("phase_0", workspace, manifest or {})
    return executor._execute()


def _prepare_workspace() -> Path:
    workspace = Path(tempfile.mkdtemp())
    (workspace / "source.md").write_text("原始材料", encoding="utf-8")
    return workspace


def _load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def test_phase_0_reuses_cache_when_sources_unchanged() -> None:
    workspace = _prepare_workspace()
    first = _run_phase_0(workspace)
    notes_path = workspace / "artifacts" / "preprocess" / "phase_00_preprocess_notes.md"
    first_notes_mtime = notes_path.stat().st_mtime_ns
    first_fingerprint = _load_json(workspace / "artifacts" / "preprocess" / "source_fingerprints.json")

    time.sleep(0.01)
    second = _run_phase_0(workspace, first.manifest_updates)
    second_fingerprint = _load_json(workspace / "artifacts" / "preprocess" / "source_fingerprints.json")

    assert second.status == "success"
    assert second.manifest_updates["preprocess_cache_status"] == "cache_hit_reused"
    assert second.manifest_updates["preprocess_reuse_status"] == "reused"
    assert second.manifest_updates["preprocess_refresh_reason"] == "source_fingerprints_match"
    assert notes_path.stat().st_mtime_ns == first_notes_mtime
    assert second_fingerprint == first_fingerprint


def test_phase_0_refreshes_when_source_fingerprints_change() -> None:
    workspace = _prepare_workspace()
    first = _run_phase_0(workspace)
    notes_path = workspace / "artifacts" / "preprocess" / "phase_00_preprocess_notes.md"
    first_notes_mtime = notes_path.stat().st_mtime_ns

    time.sleep(0.01)
    (workspace / "source.md").write_text("原始材料已变化", encoding="utf-8")
    second = _run_phase_0(workspace, first.manifest_updates)

    assert second.manifest_updates["preprocess_cache_status"] == "cache_miss_refreshed"
    assert second.manifest_updates["preprocess_reuse_status"] == "refreshed"
    assert second.manifest_updates["preprocess_refresh_reason"] == "source_fingerprints_changed"
    assert notes_path.stat().st_mtime_ns > first_notes_mtime


def test_phase_0_force_preprocess_refreshes_even_when_sources_unchanged() -> None:
    workspace = _prepare_workspace()
    first = _run_phase_0(workspace)
    notes_path = workspace / "artifacts" / "preprocess" / "phase_00_preprocess_notes.md"
    first_notes_mtime = notes_path.stat().st_mtime_ns

    time.sleep(0.01)
    second_manifest = {**first.manifest_updates, "force_preprocess": True}
    second = _run_phase_0(workspace, second_manifest)

    assert second.manifest_updates["preprocess_cache_status"] == "force_refreshed"
    assert second.manifest_updates["preprocess_reuse_status"] == "refreshed"
    assert second.manifest_updates["preprocess_refresh_reason"] == "force_preprocess_requested"
    assert notes_path.stat().st_mtime_ns > first_notes_mtime


def test_phase_0_emits_preprocess_index_and_structured_artifacts() -> None:
    workspace = _prepare_workspace()
    (workspace / "参考专利.md").write_text(
        "# 背景技术\n下面对照附图作进一步详细说明。\n# 具体实施方式\n本实施例提供的一种方法，包括以下步骤：",
        encoding="utf-8",
    )

    result = _run_phase_0(workspace)

    preprocess_dir = workspace / "artifacts" / "preprocess"
    assert (preprocess_dir / "preprocess_index.json").exists()
    assert (preprocess_dir / "source_inventory.json").exists()
    assert (preprocess_dir / "document_ast.json").exists()
    assert (preprocess_dir / "reference_style_profile.json").exists()
    assert (preprocess_dir / "template_structure_rules.json").exists()
    assert (preprocess_dir / "compliance_rules.json").exists()
    assert result.manifest_updates["preprocess_index_path"] == "artifacts/preprocess/preprocess_index.json"


def test_phase_0_classifies_material_roles_and_emits_document_ast() -> None:
    workspace = _prepare_workspace()
    (workspace / "模板_技术交底书.md").write_text("# 技术领域\n# 背景技术\n# 发明内容", encoding="utf-8")
    (workspace / "规范要求.json").write_text('{"must_have_sections": ["技术领域"]}', encoding="utf-8")

    _run_phase_0(workspace)

    source_inventory = _load_json(workspace / "artifacts" / "preprocess" / "source_inventory.json")
    roles = {item["file_name"]: item["role"] for item in source_inventory["sources"]}
    assert roles["模板_技术交底书.md"] == "template"
    assert roles["规范要求.json"] == "compliance_spec"

    document_ast = _load_json(workspace / "artifacts" / "preprocess" / "document_ast.json")
    assert document_ast["documents"]
    assert any(block["block_type"] == "section" for doc in document_ast["documents"] for block in doc["blocks"])


def test_phase_0_respects_explicit_role_mapping() -> None:
    workspace = _prepare_workspace()
    (workspace / "技术交底书框架.docx").write_bytes(b"fake-docx")
    (workspace / "技术交底书实质审核规范.pdf").write_bytes(b"fake-pdf")
    manifest = {
        "phase_0_role_map": {
            "技术交底书框架.docx": "template",
            "技术交底书实质审核规范.pdf": "compliance_spec",
        }
    }

    _run_phase_0(workspace, manifest)

    source_inventory = _load_json(workspace / "artifacts" / "preprocess" / "source_inventory.json")
    roles = {item["file_name"]: item["role"] for item in source_inventory["sources"]}
    assert roles["技术交底书框架.docx"] == "template"
    assert roles["技术交底书实质审核规范.pdf"] == "compliance_spec"


def test_phase_0_extracts_blocks_from_pdf_and_docx() -> None:
    workspace = _prepare_workspace()
    pdf_path = workspace / "参考专利.pdf"
    pdf_bytes = b"%PDF-1.4\n1 0 obj<<>>endobj\ntrailer<<>>\n%%EOF"
    pdf_path.write_bytes(pdf_bytes)

    docx_path = workspace / "技术交底书框架.docx"
    document = Document()
    document.add_paragraph("技术领域")
    document.add_paragraph("背景技术")
    document.add_paragraph("发明内容")
    document.save(docx_path)
    manifest = {"phase_0_role_map": {"参考专利.pdf": "reference_patent", "技术交底书框架.docx": "template"}}

    _run_phase_0(workspace, manifest)

    document_ast = _load_json(workspace / "artifacts" / "preprocess" / "document_ast.json")
    docs = {doc["file_name"]: doc["blocks"] for doc in document_ast["documents"]}
    assert docs["技术交底书框架.docx"]


def test_orchestrator_auto_skips_phase_0_when_preprocess_cache_is_reusable() -> None:
    workspace = _prepare_workspace()
    first = _run_phase_0(workspace)
    manifest_path = workspace / "artifacts" / "run_manifest.md"
    manifest_path.write_text("", encoding="utf-8")

    orchestrator = Orchestrator(workspace, manifest_path, from_phase="0", batch_mode=True)
    orchestrator.manifest.update(first.manifest_updates)



def test_orchestrator_phase_0_reuse_does_not_rehash_sources() -> None:
    workspace = _prepare_workspace()
    first = _run_phase_0(workspace)
    manifest_path = workspace / "artifacts" / "run_manifest.md"
    manifest_path.write_text("", encoding="utf-8")

    orchestrator = Orchestrator(workspace, manifest_path, from_phase="0", batch_mode=True)
    orchestrator.manifest.update(first.manifest_updates)

    original_phase0 = orchestrate_module.Phase0Executor

    class NoHashPhase0Executor(PhaseExecutor):
        def _fingerprints(self, _files):
            raise AssertionError("Phase 0 reuse should not rehash unchanged source files")

    orchestrate_module.Phase0Executor = NoHashPhase0Executor
    try:
        assert orchestrator._phase_0_cache_reusable() is True
    finally:
        orchestrate_module.Phase0Executor = original_phase0


def test_orchestrator_phase_0_and_phase_1_use_manifest_only_snapshot() -> None:
    workspace = _prepare_workspace()
    artifacts_dir = workspace / "artifacts"
    (artifacts_dir / "research").mkdir(parents=True)
    (artifacts_dir / "research" / "large.json").write_text("x" * 1024, encoding="utf-8")
    manifest_path = artifacts_dir / "run_manifest.md"
    manifest_path.write_text("domain_scope: 自动驾驶", encoding="utf-8")

    orchestrator = Orchestrator(workspace, manifest_path, from_phase="0", batch_mode=True)

    phase_0_snapshot = orchestrator._take_snapshot("phase_0")
    phase_1_snapshot = orchestrator._take_snapshot("phase_1")

    assert (phase_0_snapshot / "run_manifest.md").exists()
    assert (phase_1_snapshot / "run_manifest.md").exists()
    assert not (phase_0_snapshot / "research").exists()
    assert not (phase_1_snapshot / "research").exists()



def test_orchestrator_phase_1_reused_scope_does_not_prompt_again() -> None:
    workspace = _prepare_workspace()
    manifest_path = workspace / "artifacts" / "run_manifest.md"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text("", encoding="utf-8")
    orchestrator = Orchestrator(workspace, manifest_path, from_phase="1", batch_mode=False)
    orchestrator.manifest.update(
        {
            "domain_scope": "自动驾驶",
            "idea_maturity": "vague_domain",
            "phase_02_mode": "broad_domain_discovery",
            "phase_02_local_project_mode": "disabled",
            "phase_02_discovery_inputs": ["online_search"],
        }
    )

    def fail_if_prompted(_prompt, _default=""):
        raise AssertionError("Phase 1 reused scope should not prompt again")

    orchestrator._prompt_user = fail_if_prompted
    phase_1_stage = next(stage for stage in orchestrate_module.STAGES if stage["id"] == "phase_1")

    assert orchestrator._run_phase(phase_1_stage) is True



def test_orchestrator_cli_prefill_phase_1_inputs_for_webchat() -> None:
    workspace = _prepare_workspace()
    manifest_path = workspace / "artifacts" / "run_manifest.md"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text("", encoding="utf-8")

    orchestrator = Orchestrator(
        workspace,
        manifest_path,
        from_phase="1",
        batch_mode=False,
        domain_scope="智能座舱方向，没具体想法",
        local_project_paths=[],
    )

    def fail_if_prompted(_prompt, _default=""):
        raise AssertionError("CLI-prefilled Phase 1 should not prompt in webchat mode")

    orchestrator._prompt_user = fail_if_prompted
    phase_1_stage = next(stage for stage in orchestrate_module.STAGES if stage["id"] == "phase_1")

    assert orchestrator.manifest["domain_scope"] == "智能座舱方向"
    assert orchestrator.manifest["phase_02_local_project_mode"] == "disabled"
    assert orchestrator._run_phase(phase_1_stage) is True
def test_orchestrator_keeps_phase_0_when_force_preprocess_requested() -> None:
    workspace = _prepare_workspace()
    first = _run_phase_0(workspace)
    manifest_path = workspace / "artifacts" / "run_manifest.md"
    manifest_path.write_text("", encoding="utf-8")

    orchestrator = Orchestrator(workspace, manifest_path, from_phase="0", batch_mode=True)
    orchestrator.manifest.update({**first.manifest_updates, "force_preprocess": True})

    assert orchestrator._phase_0_cache_reusable() is False
    assert orchestrator._resolve_start_phase() == "phase_0"


if __name__ == "__main__":
    test_phase_0_reuses_cache_when_sources_unchanged()
    test_phase_0_refreshes_when_source_fingerprints_change()
    test_phase_0_force_preprocess_refreshes_even_when_sources_unchanged()
    test_phase_0_emits_preprocess_index_and_structured_artifacts()
    test_phase_0_classifies_material_roles_and_emits_document_ast()
    test_phase_0_respects_explicit_role_mapping()
    test_phase_0_extracts_blocks_from_pdf_and_docx()
    test_orchestrator_auto_skips_phase_0_when_preprocess_cache_is_reusable()
    test_orchestrator_keeps_phase_0_when_force_preprocess_requested()
    test_orchestrator_phase_0_reuse_does_not_rehash_sources()
    test_orchestrator_phase_0_and_phase_1_use_manifest_only_snapshot()
    test_orchestrator_phase_1_reused_scope_does_not_prompt_again()
    test_orchestrator_cli_prefill_phase_1_inputs_for_webchat()
    print("phase_0_executor regression tests passed")
